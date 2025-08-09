import os
import glob
import re
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Load environment variables
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    raise ValueError("Missing OPENAI_API_KEY in .env")
os.environ["OPENAI_API_KEY"] = openai_api_key

# Paths
PDF_FOLDER = "./pdf_documents"
DOCX_FOLDER = "./docx_documents"
TXT_FOLDER = "./txt_documents"
INDEX_PATH = "faiss_index"

def extract_text_from_pdf(path):
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            text += txt + "\n"
    return text

def extract_text_from_docx(path):
    doc = DocxDocument(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def heading_aware_chunking(text):
    # Detect headings by line pattern, e.g. ALL CAPS or numbered headings
    heading_pattern = r"(^[A-Z0-9\s\-\.,:]{3,}$)|(^\d+(\.\d+)*\s.+$)"
    lines = text.splitlines()

    chunks = []
    current_heading = "General"
    current_text = ""

    for line in lines:
        if re.match(heading_pattern, line.strip()):
            if current_text.strip():
                chunks.append(Document(
                    page_content=f"[Heading: {current_heading}]\n{current_text.strip()}",
                    metadata={"heading": current_heading}
                ))
                current_text = ""
            current_heading = line.strip()
        else:
            current_text += " " + line.strip()

    if current_text.strip():
        chunks.append(Document(
            page_content=f"[Heading: {current_heading}]\n{current_text.strip()}",
            metadata={"heading": current_heading}
        ))

    # Dynamic chunk size based on avg length
    avg_len = sum(len(d.page_content) for d in chunks) / max(1, len(chunks))
    chunk_size = max(500, min(1500, int(avg_len * 1.1)))

    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=100)
    final_chunks = []
    for doc in chunks:
        splits = splitter.split_text(doc.page_content)
        final_chunks.extend([Document(page_content=part, metadata=doc.metadata) for part in splits])

    return final_chunks

def load_all_texts():
    texts = []

    for pdf_file in glob.glob(os.path.join(PDF_FOLDER, "*.pdf")):
        print(f"Loading PDF: {pdf_file}")
        texts.append(extract_text_from_pdf(pdf_file))

    for docx_file in glob.glob(os.path.join(DOCX_FOLDER, "*.docx")):
        print(f"Loading DOCX: {docx_file}")
        texts.append(extract_text_from_docx(docx_file))

    for txt_file in glob.glob(os.path.join(TXT_FOLDER, "*.txt")):
        print(f"Loading TXT: {txt_file}")
        texts.append(extract_text_from_txt(txt_file))

    return texts

def main():
    texts = load_all_texts()
    if not texts:
        print("No documents found in folders.")
        return

    print("Performing heading-aware chunking...")
    all_chunks = []
    for text in texts:
        chunks = heading_aware_chunking(text)
        all_chunks.extend(chunks)

    print(f"Total chunks created: {len(all_chunks)}")

    print("Creating OpenAI embeddings and building FAISS index...")
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_documents(all_chunks, embeddings)

    vectorstore.save_local(INDEX_PATH)
    print(f"FAISS index saved at '{INDEX_PATH}'")

if __name__ == "__main__":
    main()
